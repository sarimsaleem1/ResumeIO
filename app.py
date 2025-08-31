import streamlit as st
import requests
from bs4 import BeautifulSoup
import io
from xhtml2pdf import pisa
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import time

# ApplicationRunner class moved directly into app.py
class ApplicationRunner:
    def __init__(self):
        self.base_url = "https://resume.io/api/app/resumes"
        self.headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36"
        }
    
    def get_resume_html(self, token):
        """Fetch resume HTML from resume.io API"""
        try:
            response = requests.get(f"{self.base_url}/{token}", headers=self.headers)
            response.raise_for_status()
            data = response.json()
            return data.get("html", "")
        except Exception as e:
            print(f"Error fetching resume: {str(e)}")
            return None
    
    def html_to_pdf(self, html_content):
        """Convert HTML content to PDF format"""
        try:
            pdf_bytes = io.BytesIO()
            pisa.CreatePDF(io.StringIO(html_content), dest=pdf_bytes)
            pdf_bytes.seek(0)
            return pdf_bytes.getvalue()
        except Exception as e:
            print(f"Error generating PDF: {str(e)}")
            return None
    
    def html_to_docx(self, html_content):
        """Convert HTML content to DOCX format"""
        try:
            doc = Document()
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Process HTML elements
            for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'ul', 'ol', 'li', 'strong', 'em', 'u', 'a']):
                if element.name.startswith('h'):
                    # Handle headings
                    level = int(element.name[1])
                    heading = doc.add_heading('', level=level)
                    self._process_text_element(element, heading)
                
                elif element.name == 'p':
                    # Handle paragraphs
                    paragraph = doc.add_paragraph()
                    self._process_text_element(element, paragraph)
                
                elif element.name in ['ul', 'ol']:
                    # Handle lists
                    list_type = 'unordered' if element.name == 'ul' else 'ordered'
                    for li in element.find_all('li', recursive=False):
                        paragraph = doc.add_paragraph(
                            style='List Bullet' if list_type == 'unordered' else 'List Number'
                        )
                        self._process_text_element(li, paragraph)
            
            # Save to bytes
            doc_bytes = io.BytesIO()
            doc.save(doc_bytes)
            doc_bytes.seek(0)
            return doc_bytes.getvalue()
        except Exception as e:
            print(f"Error generating DOCX: {str(e)}")
            return None
    
    def _process_text_element(self, element, doc_element):
        """Process text elements with formatting"""
        for child in element.children:
            if child.name is None:
                # Plain text
                run = doc_element.add_run(str(child))
            elif child.name == 'strong':
                # Bold text
                run = doc_element.add_run(child.get_text())
                run.bold = True
            elif child.name == 'em':
                # Italic text
                run = doc_element.add_run(child.get_text())
                run.italic = True
            elif child.name == 'u':
                # Underlined text
                run = doc_element.add_run(child.get_text())
                run.underline = True
            elif child.name == 'a':
                # Hyperlinks
                run = doc_element.add_run(child.get_text())
                run.underline = True
                # Note: Adding actual hyperlinks requires more complex handling
            else:
                # Other elements - process recursively
                self._process_text_element(child, doc_element)

# Page configuration
st.set_page_config(
    layout="wide", 
    page_title="Resume IO Downloader"
)

# Session state initialization
def init_session_state():
    if "render_token" not in st.session_state:
        st.session_state["render_token"] = ""
    if "render_token_bool" not in st.session_state:
        st.session_state["render_token_bool"] = True
    if "current_page" not in st.session_state:
        st.session_state["current_page"] = "Home"

init_session_state()

# Navigation handler
def navigate_to(page):
    st.session_state["current_page"] = page

# Navigation
st.sidebar.title("Navigation")
if st.sidebar.button("Home", use_container_width=True):
    navigate_to("Home")
if st.sidebar.button("Download Resume", use_container_width=True):
    navigate_to("Download Resume")
if st.sidebar.button("About", use_container_width=True):
    navigate_to("About")

# Highlight current page in sidebar
if st.session_state["current_page"] == "Home":
    st.sidebar.markdown("**Home**")
elif st.session_state["current_page"] == "Download Resume":
    st.sidebar.markdown("**Download Resume**")
elif st.session_state["current_page"] == "About":
    st.sidebar.markdown("**About**")

# Home Page
if st.session_state["current_page"] == "Home":
    st.title("Resume IO Downloader")
    st.write("""
        To fetch your resume from resume.io, you need to get your `rendering token` from 
        https://resume.io/api/app/resumes
    """)
    
    st.markdown("""
    ### How to get your renderToken:
    1. Go to https://resume.io and open your resume
    2. Open your browser's developer tools (F12)
    3. Go to the Network tab
    4. Refresh the page and look for a request to `https://resume.io/api/app/resumes/[some_id]`
    5. Click on that request and check the Response tab
    6. Copy the `rendering token` from the response
    """)
    
    render_token = st.text_input("Enter renderToken", placeholder="Paste renderToken here...")
    st.session_state["render_token"] = render_token
    
    if st.button("Submit"):
        if st.session_state["render_token"]:
            st.session_state["render_token_bool"] = False
            st.success("Token accepted! Proceed to download page.")
            # Automatically navigate to download page after successful token submission
            navigate_to("Download Resume")
        else:
            st.error("renderToken required to proceed.")
            st.session_state["render_token_bool"] = True
    
    st.button(
        "Go to Download Resume", 
        disabled=st.session_state["render_token_bool"],
        on_click=lambda: navigate_to("Download Resume")
    )

# Download Resume Page
elif st.session_state["current_page"] == "Download Resume":
    if st.session_state["render_token_bool"]:
        st.error("renderToken required to proceed.")
        st.button("Back to Home", on_click=lambda: navigate_to("Home"))
    else:
        st.title("Download Your Resume")
        
        # Display the token for verification (you can remove this line if you want)
        st.write(f"Token: {st.session_state['render_token']}")
        
        try:
            # Show loading animation
            with st.spinner("Creating your document..."):
                time.sleep(1)  # Simulate processing
                
                # Create an instance of ApplicationRunner
                app_runner = ApplicationRunner()
                
                # Get HTML content
                token = str(st.session_state["render_token"]).strip()
                html_content = app_runner.get_resume_html(token)
                
                # Debug information (you can remove these lines if you want)
                if html_content is None:
                    st.error("Failed to fetch resume. Please check your token.")
                    st.button("Try Again", on_click=lambda: navigate_to("Home"))
                else:
                    # Generate PDF and DOCX
                    pdf_bytes = app_runner.html_to_pdf(html_content)
                    docx_bytes = app_runner.html_to_docx(html_content)
                    
                    if pdf_bytes is None and docx_bytes is None:
                        st.error("Failed to generate documents. Please try again.")
                        st.button("Try Again", on_click=lambda: navigate_to("Home"))
                    else:
                        # Download buttons
                        col1, col2 = st.columns(2)
                        
                        with col1:
                            if pdf_bytes is not None:
                                st.download_button(
                                    label="📄 Download PDF",
                                    data=pdf_bytes,
                                    file_name=f"resume_{token}.pdf",
                                    mime="application/pdf"
                                )
                        
                        with col2:
                            if docx_bytes is not None:
                                st.download_button(
                                    label="📝 Download DOCX",
                                    data=docx_bytes,
                                    file_name=f"resume_{token}.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                                )
                        
                        st.success("Resume generated successfully!")
                        
                        # Add option to download another resume
                        st.button("Download Another Resume", on_click=lambda: navigate_to("Home"))
                        
        except Exception as e:
            st.error(f"Error: {str(e)}")
            st.button("Try Again", on_click=lambda: navigate_to("Home"))

# About Page
elif st.session_state["current_page"] == "About":
    st.title("About")
    st.info("This app allows you to download your resume from resume.io in both PDF and DOCX formats.")
    st.button("Back to Home", on_click=lambda: navigate_to("Home"))
